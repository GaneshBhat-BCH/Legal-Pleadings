import os
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- 1. Mermaid to Image Conversion ---
def extract_mermaid_code(markdown_text: str) -> list[str]:
    """Extracts all mermaid code blocks from the markdown."""
    diagrams = []
    lines = markdown_text.split('\n')
    in_block = False
    current_block = []
    
    for line in lines:
        if line.strip().startswith('```mermaid'):
            in_block = True
            current_block = []
        elif line.strip() == '```' and in_block:
            in_block = False
            diagrams.append('\n'.join(current_block))
        elif in_block:
            current_block.append(line)
            
    return diagrams

def generate_mermaid_image(mermaid_code: str, output_path: str):
    """Uses Kroki API to render Mermaid code as a PNG."""
    import base64
    import zlib
    
    # Kroki expects compressed and base64 encoded payload
    compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
    
    url = f"https://kroki.io/mermaid/png/{encoded}"
    
    # Optional styling for higher contrast
    headers = {'Accept': 'image/png'}
    
    print(f"Generating diagram: {output_path}")
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        with open(output_path, 'wb') as f:
            f.write(response.content)
    else:
        print(f"Failed to generate diagram: {response.text}")

# --- 2. Word Document Generation ---
def create_sdd_docx(md_path: str, output_docx: str):
    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Create Document
    doc = Document()
    
    # Styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Solution Design Document (SDD)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Legal Pleadings RAG & Processing Engine')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Extract Diagrams
    diagrams = extract_mermaid_code(md_text)
    image_paths = []
    
    # Generate Images temporarily
    for i, code in enumerate(diagrams):
        img_path = f"temp_diagram_{i}.png"
        generate_mermaid_image(code, img_path)
        image_paths.append(img_path)

    # Parse and add content (simplified parser for this specific SDD)
    lines = md_text.split('\n')
    diagram_index = 0
    in_code_block = False
    code_content = []

    for line in lines:
        if line.startswith('```mermaid'):
            in_code_block = True
            
            # Insert the generated image
            if diagram_index < len(image_paths) and os.path.exists(image_paths[diagram_index]):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(image_paths[diagram_index], width=Inches(6.0))
                
                caption = doc.add_paragraph(f"Figure {diagram_index + 1}: Architecture Diagram")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.style = doc.styles['Caption']
                
            diagram_index += 1
            continue
            
        elif line.startswith('```text') or line.startswith('```json'):
            in_code_block = True
            code_content = []
            continue
            
        elif line.startswith('```') and in_code_block:
            in_code_block = False
            # Add code block
            p = doc.add_paragraph('\n'.join(code_content))
            p.style = doc.styles['No Spacing']
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            continue
            
        if in_code_block:
            code_content.append(line)
            continue

        # Normal Text Processing
        if not line.strip():
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            # Basic bolding `**text**` inline parsing could be done, but keeping it simple
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)

    # Save
    doc.save(output_docx)
    print(f"\nSuccessfully created Word Document: {output_docx}")

    # Cleanup temp images
    for img in image_paths:
        if os.path.exists(img):
            os.remove(img)

if __name__ == "__main__":
    create_sdd_docx("Legal_Pleadings_SDD.md", "Legal_Pleadings_SDD.docx")

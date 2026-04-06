import sys
from docx import Document
import argparse

def read_docx(file_path, start_line, end_line):
    try:
        doc = Document(file_path)
        print(f"--- PARAGRAPHS ({start_line} to {end_line}) ---")
        for i, paragraph in enumerate(doc.paragraphs):
            if i < start_line - 1: continue
            if i >= end_line: break
            print(f"Para {i+1}: {paragraph.text}")
        
        print("\n--- TABLES (FIRST 5) ---")
        for i, table in enumerate(doc.tables):
            if i >= 5: break
            print(f"\nTable {i}:")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.replace('\n', ' ').strip())
                print(" | ".join(row_text))
        
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="Path to docx file")
    parser.add_argument("--start", type=int, default=1, help="Start paragraph (1-indexed)")
    parser.add_argument("--end", type=int, default=50, help="End paragraph (1-indexed)")
    args = parser.parse_args()
    
    read_docx(args.path, args.start, args.end)

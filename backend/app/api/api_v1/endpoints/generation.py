import os
from pathlib import Path
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field
from typing import List, Dict, Any
import requests
import json
import time
from docx import Document
from app.core.config import settings
from app.services.rag_service import retrieve_documents, vector_store
from app.core.logger import activity_logger

router = APIRouter()

class AllegationPoint(BaseModel):
    point_number: int
    allegation_text: str
    is_rebuttable: bool
    user_response: str = Field(default="", description="The response or additional context provided by the user via Co-Pilot")

class DocumentMetadata(BaseModel):
    charging_party: str
    respondent: str
    date_filed: str
    legal_case_summary: str
    all_detected_categories: str

class GenerationRequest(BaseModel):
    document_metadata: DocumentMetadata
    allegations_list: List[AllegationPoint]
    allegation_classification: List[Dict[str, Any]] = []
    defense_and_proofs: List[Dict[str, Any]] = []

@router.post("/generate_statement")
async def generate_statement(request: GenerationRequest):
    party_name = request.document_metadata.charging_party
    activity_logger.log_event("Generation", "START", party_name, f"Generating Position Statement based on {len(request.allegations_list)} allegations.")
    
    api_key = settings.AZURE_OPENAI_API_KEY
    endpoint = settings.AZURE_OPENAI_ENDPOINT
    deployment_name = "gpt-5"
    
    # 1. Prepare RAG Context
    # We aggregate text from the allegations to query the vector store
    search_query = request.document_metadata.legal_case_summary
    for pt in request.allegations_list:
        search_query += f" {pt.allegation_text} {pt.user_response}"
        
    try:
        # We query the LangChain vector store directly via rag_service
        # retrieve_documents expects query, k
        rag_docs = await retrieve_documents(query=search_query, k=5)
        rag_context = "\n\n".join([d.page_content for d in rag_docs])
    except Exception as e:
        print(f"Warning: RAG retrieval failed: {e}")
        rag_context = "No additional context found."

    # 2. Build Prompt for Azure OpenAI
    # We instruct the LLM to draft a formal Position Statement
    system_prompt = """You are a Senior Legal Counsel drafting a formal Position Statement on behalf of a Respondent employer.
Your goal is to consume the structured facts (Allegations and User Responses) and the retrieved Legal Citations (RAG Context) to produce a cohesive, professional, and robust Position Statement suitable for submission to an agency (e.g., EEOC).

Structure the Position Statement as follows:
1. Introduction: State the parties and summarize the respondent's definitive stance (using User Responses).
2. Legal Framework: Interweave the provided RAG Context laws gracefully into the defense.
3. Statement of Facts & Rebuttal: Address each allegation point-by-point, applying the user's response to refute or contextualize the claim.
4. Conclusion: Summarize the defense and unequivocally request dismissal.

[CONTENT FILTER & COMPLIANCE RULES]
To comply with strict Azure OpenAI content safety filters, you MUST NOT generate explicitly sexual, violent, or hate-related language in your draft. Instead, use clinical, sterile, and professional legal terminology to summarize any sensitive facts (e.g., use "alleged inappropriate physical contact" instead of explicit descriptions). Ensure the final document is completely sanitized of severe explicit language while remaining legally persuasive.

Ensure the tone is extremely formal, persuasive, and legally sound. Do not use markdown. Return only the raw text of the document."""

    user_prompt = f"""
[DOCUMENT METADATA]
Charging Party: {request.document_metadata.charging_party}
Respondent: {request.document_metadata.respondent}
Date Filed: {request.document_metadata.date_filed}
Summary: {request.document_metadata.legal_case_summary}
Categories: {request.document_metadata.all_detected_categories}

[ALLEGATIONS AND USER RESPONSES]
"""
    for pt in request.allegations_list:
        user_prompt += f"\nPoint {pt.point_number}:\nAllegation: {pt.allegation_text}\nUser Response: {pt.user_response}\n"

    user_prompt += f"\n\n[RELEVANT LAW VIA RAG]\n{rag_context}\n"

    # Call Azure OpenAI Chat Completions using exact endpoint
    if "chat/completions" in settings.AZURE_OPENAI_ENDPOINT:
        chat_url = settings.AZURE_OPENAI_ENDPOINT
    else:
        chat_url = f"{settings.AZURE_OPENAI_ENDPOINT.rstrip('/')}/openai/deployments/{deployment_name}/chat/completions?api-version=2024-05-01-preview"
    
    headers = {
        "api-key": api_key,
        "Content-Type": "application/json"
    }
    payload = {
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
    }

    retry_delay = 5
    attempt = 0
    while True:
        try:
            # Standardized timeout to 1200s (20 minutes)
            chat_res = requests.post(chat_url, headers=headers, json=payload, timeout=1200)
            if chat_res.status_code != 200:
                if chat_res.status_code in [401, 403]:
                    err_msg = f"Chat generation failed due to credentials: {chat_res.status_code} - {chat_res.text}"
                    activity_logger.log_event("Generation", "ERROR", party_name, err_msg)
                    raise HTTPException(status_code=401, detail="Authentication failed. Please check AI credentials.")
                    
                err_msg = f"Chat generation failed: {chat_res.status_code} - {chat_res.text}"
                activity_logger.log_event("Generation", "ERROR", party_name, err_msg)
                raise Exception(err_msg)
            
            completion = chat_res.json()
            generated_text = completion["choices"][0]["message"]["content"]
            break # Loop break on success
        except HTTPException as e:
            raise e
        except Exception as e:
            err_msg = str(e)
            activity_logger.log_ai_error(err_msg)
            attempt += 1
            activity_logger.log_event("Generation", "WARNING", party_name, f"Retrying infinite loop (attempt {attempt})...")
            time.sleep(retry_delay)
            continue

    # 3. Create Word Document
    try:
        doc = Document()
        doc.add_heading('POSITION STATEMENT', 0)
        
        # Add a bit of metadata at the top
        doc.add_paragraph(f"Charging Party: {request.document_metadata.charging_party}")
        doc.add_paragraph(f"Respondent: {request.document_metadata.respondent}")
        doc.add_paragraph(f"Date: {request.document_metadata.date_filed}")
        doc.add_paragraph("_" * 40)
        
        # Add generated content
        for paragraph in generated_text.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        # 4. Save to Downloads folder
        downloads_path = Path.home() / "Downloads"
        file_name = f"Position_Statement_{request.document_metadata.charging_party.replace(' ', '_')}.docx"
        file_path = downloads_path / file_name
        
        doc.save(str(file_path))
        activity_logger.log_event("Generation", "SUCCESS", party_name, f"Position statement saved to {file_path}")
    except Exception as e:
        err_msg = f"Failed to generate Word document: {e}"
        activity_logger.log_event("Generation", "ERROR", party_name, err_msg)
        raise HTTPException(status_code=500, detail=err_msg)

    return {
        "status": "success",
        "message": "Position statement generated successfully.",
        "file_path": str(file_path)
    }

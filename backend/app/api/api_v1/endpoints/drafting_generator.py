from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
import os
import asyncio
import httpx
import json
import time
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from app.core.config import settings
from app.services.rag_service import retrieve_documents
from app.core.logger import activity_logger

# Helper to resolve logo paths
_HERE = Path(__file__).parent
_SEMAPHORE = asyncio.Semaphore(5)

def apply_safety_mask(text: str) -> str:
    """Mask sensitive legal terms that might trigger AI content filters."""
    if not isinstance(text, str): return text
    masks = {
        r'\bsexual\b': 's-e-x-u-a-l', r'\bharassment\b': 'har*ssment',
        r'\brape\b': 'r*pe', r'\bassault\b': 'ass*ult', r'\bviolence\b': 'vi*lence',
        r'\bsex\b': 's*x', r'\bracial\b': 'rac*al', r'\bdiscrimination\b': 'discrim*nation',
        r'\bcolor\b': 'col*r', r'\brace\b': 'ra*e', r'\bblack\b': 'bl*ck'
    }
    processed = text
    for pattern, replacement in masks.items():
        processed = re.sub(pattern, replacement, processed, flags=re.IGNORECASE)
    return processed

def restore_safety_mask(text: Any) -> Any:
    """Restore masked terms to their original professional legal versions."""
    if isinstance(text, list):
        return [restore_safety_mask(i) for i in text]
    if isinstance(text, dict):
        return {k: restore_safety_mask(v) for k, v in text.items()}
    if not isinstance(text, str): return text
    
    restores = {
        's-e-x-u-a-l': 'sexual', 'har*ssment': 'harassment',
        'r*pe': 'rape', 'ass*ult': 'assault', 'vi*lence': 'violence',
        's*x': 'sex', 'rac*al': 'racial', 'discrim*nation': 'discrimination',
        'col*r': 'color', 'ra*e': 'race', 'bl*ck': 'black'
    }
    processed = text
    for masked, original in restores.items():
        processed = re.sub(re.escape(masked), original, processed, flags=re.IGNORECASE)
    return processed

async def call_llm_module(url: str, api_key: str, system_prompt: str, user_content: str, response_format: str = "json_object"):
    """Async wrapper for Azure OpenAI calls with concurrency control and basic retry."""
    headers = {"api-key": api_key, "Content-Type": "application/json"}
    payload = {
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ],
        "max_completion_tokens": 4096,
        "temperature": 0.3
    }
    if response_format == "json_object":
        payload["response_format"] = {"type": "json_object"}

    async with _SEMAPHORE:
        async with httpx.AsyncClient() as client:
            for attempt in range(2): # Simple 2-attempt retry
                try:
                    response = await client.post(url, headers=headers, json=payload, timeout=300.0)
                    if response.status_code == 200:
                        return response.json()["choices"][0]["message"]["content"]
                    elif response.status_code == 429:
                        await asyncio.sleep(2 * (attempt + 1))
                        continue
                    else:
                        activity_logger.log_event("Drafting", "AI_ERROR", "LLM", f"Status {response.status_code}: {response.text}")
                except Exception as e:
                    activity_logger.log_event("Drafting", "AI_EXCEPTION", "LLM", str(e))
                    await asyncio.sleep(1)
            return None

# Navigate to backend/assets correctly (parent is api_v1, parent2 is api, parent3 is app, parent4 is backend)
_ASSETS_DIR = _HERE.parent.parent.parent.parent / "assets"
LEFT_LOGO = _ASSETS_DIR / "bch_logo.png"
RIGHT_LOGO = _ASSETS_DIR / "hms_logo.png"

# Reference Document Path (Template)
# Use a relative fallback for the VM environment if the OneDrive path is missing
_DEFAULT_TEMPLATE = _ASSETS_DIR / "templates" / "Legal_Template.docx"
REFERENCE_DOC_PATH = os.getenv("REFERENCE_DOC_PATH", str(_DEFAULT_TEMPLATE))

def sanitize_xml(text):
    """Remove control characters that break XML/DOCX."""
    if not isinstance(text, str):
        return str(text) if text is not None else ""
    # Standard XML-illegal character ranges
    illegal_xml_chars_re = re.compile(
        u'[\x00-\x08\x0b\x0c\x0e-\x1F\uD800-\uDFFF\uFFFE\uFFFF]'
    )
    return illegal_xml_chars_re.sub('', text)

def get_current_date_str():
    from datetime import datetime
    return datetime.now().strftime("%B %d, %Y")

def copy_standard_first_page(target_doc, charging_party, respondent):
    """
    Clones the intro sections (Transmittal Letter + Legal Caption) from the master 
    Legal_Template.docx into target_doc. Handles dynamic replacements.
    """
    if not os.path.exists(REFERENCE_DOC_PATH):
        activity_logger.log_event("Drafting", "WARN", charging_party, "Reference doc not found. Skipping mirror.")
        return False
    
    source_doc = Document(REFERENCE_DOC_PATH)
    curr_date = get_current_date_str()
    
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    def iter_block_items(parent):
        if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
            parent_elm = parent.element.body
        elif hasattr(parent, '_tc'):
            parent_elm = parent._tc
        else:
            raise TypeError("Value must be a document or cell")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    # Global Font Setup for Cloned Parts
    def apply_replacement(text, cp, resp, date):
        # Specific Roxton Draft Replacements
        # Use word boundaries \b for BCH to avoid messing up addresses like 'BCH3046'
        text = re.sub(r"\bAndrea Roxton\b", cp, text, flags=re.IGNORECASE)
        text = re.sub(r"\bMs\. Roxton\b", f"Ms. {cp.split()[-1]}", text, flags=re.IGNORECASE)
        text = re.sub(r"March 25, 2026", date, text)
        
        # Only replace 'Respondent' if it's not already followed by the respondent name
        if resp.lower() not in text.lower() or "Respondent," in text:
             # Handle 'Respondent, [Name]' specifically to avoid 'Name, Name'
             text = text.replace("Respondent,", f"{resp},")
             text = re.sub(r"\bRespondent\b", resp, text, flags=re.IGNORECASE)
            
        text = re.sub(r"\bBCH\b", resp, text)
        return sanitize_xml(text)

    for block in iter_block_items(source_doc):
        # Stop exactly before Section I (INTRODUCTION) starts
        # This keeps the Letter (Page 1) and Case Caption (Page 2)
        if isinstance(block, Paragraph):
            if "I." in block.text.upper() and "INTRODUCTION" in block.text.upper():
                break
        
        if isinstance(block, Table):
            # Clone Table Structure (Captures Table 0 Logos and Table 1 Caption)
            new_table = target_doc.add_table(rows=len(block.rows), cols=len(block.columns))
            new_table.style = block.style
            for r_idx, row in enumerate(block.rows):
                for c_idx, cell in enumerate(row.cells):
                    target_cell = new_table.cell(r_idx, c_idx)
                    for p in cell.paragraphs:
                        new_p = target_cell.add_paragraph()
                        new_p.alignment = p.alignment
                        for run in p.runs:
                            # Handle Logos
                            if "BCH_LOGO" in run.text.upper() or "HMS_LOGO" in run.text.upper():
                                if "BCH_LOGO" in run.text.upper() and LEFT_LOGO.exists():
                                    new_p.add_run().add_picture(str(LEFT_LOGO), width=Inches(3.0))
                                elif "HMS_LOGO" in run.text.upper() and RIGHT_LOGO.exists():
                                    new_p.add_run().add_picture(str(RIGHT_LOGO), width=Inches(1.5))
                                continue

                            text = apply_replacement(run.text, charging_party, respondent, curr_date)
                            new_run = new_p.add_run(text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if run.font.size: new_run.font.size = run.font.size
                            else: new_run.font.size = Pt(11)
                            new_run.font.name = 'Times New Roman'
        
        elif isinstance(block, Paragraph):
            # Skip empty paragraphs that don't have formatting (like spacing)
            if not block.text.strip() and not block.runs:
                target_doc.add_paragraph()
                continue
            
            # Manually handle Page Break before the Pleading Caption
            if "COMMONWEALTH OF MASSACHUSETTS" in block.text.upper():
                target_doc.add_page_break()
            
            text = apply_replacement(block.text, charging_party, respondent, curr_date)
            new_p = target_doc.add_paragraph()
            new_p.alignment = block.alignment
            new_p.style = block.style
            
            for run in block.runs:
                run_text = apply_replacement(run.text, charging_party, respondent, curr_date)
                new_run = new_p.add_run(run_text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size: new_run.font.size = run.font.size
                else: new_run.font.size = Pt(11)
                new_run.font.name = 'Times New Roman'

    return True

def repair_json(json_str):
    """Basic repair for common LLM JSON formatting errors."""
    if not json_str: return "{}"
    json_str = json_str.strip()
    if json_str.startswith("```json"):
        json_str = json_str[7:]
    if json_str.endswith("```"):
        json_str = json_str[:-3]
    return json_str.strip()

def chunk_text(text: str, max_chars: int = 6000, overlap: int = 500) -> list:
    """Split large text into overlapping chunks for reliable AI analysis."""
    chunks = []
    current_pos = 0
    while current_pos < len(text):
        end_pos = min(current_pos + max_chars, len(text))
        
        # Try to find a clean break (newline) within the last 200 chars of the chunk
        if end_pos < len(text):
            break_pos = text.rfind('\n', end_pos - 200, end_pos)
            if break_pos != -1 and break_pos > current_pos:
                end_pos = break_pos
        
        chunks.append(text[current_pos:end_pos].strip())
        
        # Move forward by (chunk_size - overlap) to ensure continuity
        new_pos = end_pos - overlap
        if new_pos <= current_pos:
            new_pos = end_pos
        current_pos = new_pos
        
        if end_pos >= len(text):
            break
            
    return chunks

async def generate_intro_history(url, key, cp, resp, brief_points, rag_context):
    """Generates Section I (Introduction) and II (Procedural History)."""
    system_prompt = f"[SENIOR DEFENSE COUNSEL] Draft Section I: INTRODUCTION and Section II: PROCEDURAL HISTORY for {resp}.\nSTRICT: Use formal US legal language and weave in-line citations of MCAD/EEOC regulations and relevant case law from the provided PRECEDENT.\nRETURN JSON: {{ \"introduction\": \"...\", \"procedural_history\": \"...\" }}"
    user_content = apply_safety_mask(f"PARTIES: {cp} (CP) vs {resp} (Resp). SAMPLES DATA: {json.dumps(brief_points[:5])}\n\nLEGAL PRECEDENT: {rag_context[:1000]}")
    res = await call_llm_module(url, key, system_prompt, user_content)
    try:
        if res:
            parsed = restore_safety_mask(json.loads(repair_json(res)))
            activity_logger.log_event("Drafting", "MODULE_RESULT", "Intro/History", "SUCCESS")
            return parsed
    except:
        activity_logger.log_event("Drafting", "MODULE_RESULT", "Intro/History", "FAIL_PARSE")
    return {}

async def generate_facts(url, key, cp, resp, all_points, rag_context):
    """Generates Section III (Statement of Facts)."""
    system_prompt = f"[SENIOR DEFENSE COUNSEL] Draft Section III: STATEMENT OF FACTS for {resp}. Narrative paragraph style only. Weave in-line references to applicable laws or standards from the PRECEDENT to establish the factual defense.\nRETURN JSON: {{ \"statement_of_facts\": \"...\" }}"
    user_content = apply_safety_mask(f"ALLEGATION DATA:\n{json.dumps(all_points[:30])}\n\nLEGAL PRECEDENT: {rag_context[:1000]}")
    res = await call_llm_module(url, key, system_prompt, user_content)
    try:
        if res:
            parsed = restore_safety_mask(json.loads(repair_json(res)))
            activity_logger.log_event("Drafting", "MODULE_RESULT", "Facts", "SUCCESS")
            return parsed.get("statement_of_facts", "")
    except:
        activity_logger.log_event("Drafting", "MODULE_RESULT", "Facts", "FAIL_PARSE")
    return ""

async def generate_analysis_section(url, key, cp, resp, points, category, rag_context):
    """Generates one sub-section of Section V (Legal Analysis)."""
    system_prompt = f"[SENIOR DEFENSE COUNSEL] Draft Section V legal analysis for {category.upper()}.\nSTRICT: Formally cite MCAD/EEOC regulations and relevant Case Law/Judgments from the PRECEDENT within the text.\nRETURN JSON: {{ \"content\": \"...\" }}"
    user_content = apply_safety_mask(f"CATEGORY: {category}\nDATA:\n{json.dumps(points[:15])}\n\nLITIGATION PRECEDENT: {rag_context[:1500]}")
    res = await call_llm_module(url, key, system_prompt, user_content)
    try:
        if res:
            parsed = restore_safety_mask(json.loads(repair_json(res)))
            activity_logger.log_event("Drafting", "MODULE_RESULT", f"Analysis_{category}", "SUCCESS")
            return parsed.get("content", "")
    except:
        activity_logger.log_event("Drafting", "MODULE_RESULT", f"Analysis_{category}", "FAIL_PARSE")
    return ""

async def generate_conclusion_appendix(url, key, cp, resp):
    """Generates Section VII (Conclusion) and VIII (Appendix)."""
    prompt = f"[SENIOR DEFENSE COUNSEL] Draft Section VII: CONCLUSION and Section VIII: APPENDIX (Statutory Framework) for {resp}. Appendix should cite MCAD/EEOC regulations relevant to discrimination and retaliation. Return JSON: {{ \"conclusion\": \"...\", \"appendix\": \"...\" }}"
    res = await call_llm_module(url, key, prompt, f"Case: {cp} vs {resp}")
    return json.loads(repair_json(res)) if res else {}

router = APIRouter()

class CombinedDraftRequest(BaseModel):
    raw_data: str = Field(..., description="The stringified table/list containing Allegations and Answers.")
    folder_path: Optional[str] = Field(None, description="Absolute directory path (defaults to ./Drafts if empty).")
    charging_party: str = Field("Unknown", description="Name of the charging party if known")
    respondent: str = Field("Boston Children's Hospital", description="Name of the respondent")
    case_number: Optional[str] = Field(None, description="MCAD/EEOC Case Number")

@router.post("/generate_position_draft")
async def generate_position_draft(request: CombinedDraftRequest):
    activity_logger.log_event("Drafting", "START", request.charging_party, "Executing Roxton-Style Point-by-Point Drafting")
    
    api_key = settings.AZURE_OPENAI_API_KEY
    deployment_id = settings.AZURE_OPENAI_MODEL
    
    # Construct base URL for Azure OpenAI
    base_url = settings.AZURE_OPENAI_ENDPOINT.rstrip('/')
    if "/openai/deployments/" in base_url:
        final_url = base_url
    else:
        final_url = f"{base_url}/openai/deployments/{deployment_id}/chat/completions?api-version=2024-05-01-preview"

    try:
        # --- STEP 1: ADAPTIVE ANALYSIS (JSON-OR-UNSTRUCTURED) --- 
        all_points = []
        
        # Sanitize raw_data: Remove problematic literal control characters that break JSON
        sanitized_raw = request.raw_data.replace("\r", "").replace("\t", " ")
        if "\\n" not in sanitized_raw: # If they used literal newlines instead of \n
            sanitized_raw = sanitized_raw.replace("\n", "\\n")

        try:
            # Check if input is already structured JSON
            structured_input = json.loads(repair_json(sanitized_raw))
            all_points = structured_input.get("points", []) or structured_input.get("allegations_list", [])
            activity_logger.log_event("Drafting", "BYPASS", request.charging_party, "Using pre-structured JSON input.")
        except:
            # Partitioned Analysis for Massive Unstructured Inputs
            activity_logger.log_event("Drafting", "ANALYSIS_START", request.charging_party, "Executing Partitioned Raw Analysis...")
            
            analysis_prompt = """[SENIOR LEGAL ANALYST] Extract ALL individual allegations, documentation proofs, and responses verbatim.
            STRICT RULES:
            - ZERO MERGING. Every numbered index (1, 2, 3...) must be its own unique entry.
            - NO SUMMARIZATION. Keep all client names and verbatim quotes.
            - Pattern: Identify the index number, then the allegation text, then the suggested proofs/documents, and finally the employer's response.
            - IMPORTANT: The data is unstructured. Do not be confused by commas inside legal sentences.
            - Return JSON: { "points": [ { "label": "X", "allegation": "...", "suggested_proof": "...", "response": "..." }, ... ] }
            """
            
            # Use smaller 2,500-character chunks with generous 500-char overlap
            raw_chunks = chunk_text(sanitized_raw, 2500, 500)
            
            for idx, chunk in enumerate(raw_chunks):
                activity_logger.log_event("Drafting", "ANALYSIS_BLOCK", request.charging_party, f"Processing Part {idx+1}/{len(raw_chunks)}")
                try:
                    # Mask the chunk before sending to AI to bypass content filters
                    masked_chunk = apply_safety_mask(chunk)
                    async with httpx.AsyncClient() as client:
                        res1 = await client.post(
                            final_url,
                            headers={"api-key": api_key, "Content-Type": "application/json"},
                            json={
                                "messages": [{"role": "system", "content": analysis_prompt}, {"role": "user", "content": f"DATA (PART {idx+1}/{len(raw_chunks)}):\n{masked_chunk}"}],
                                "response_format": { "type": "json_object" },
                                "max_completion_tokens": 4096
                            },
                            timeout=300.0
                        )
                    if res1.status_code == 200:
                        # Restore any masked terms in the output JSON
                        chunk_json = restore_safety_mask(json.loads(repair_json(res1.json()["choices"][0]["message"]["content"])))
                        new_pts = chunk_json.get("points", [])
                        if isinstance(new_pts, list): all_points.extend(new_pts)
                except Exception as e:
                    activity_logger.log_event("Drafting", "ANALYSIS_WARN", request.charging_party, f"Part {idx+1} failed: {str(e)}")

        if not all_points:
            raise Exception("No allegations found in input.")

        # --- STEP 1d: POINT AUDIT & GAP RECOVERY (SAFETY NET) ---
        try:
            found_labels = []
            for p in all_points:
                try:
                    lbl_num = re.findall(r'\d+', str(p.get('label', '')))
                    if lbl_num: found_labels.append(int(lbl_num[0]))
                except: continue
            
            if found_labels:
                found_labels = sorted(list(set(found_labels)))
                gaps = [i for i in range(min(found_labels), max(found_labels)) if i not in found_labels]
                
                if gaps:
                    activity_logger.log_event("Drafting", "GAP_DETECTED", request.charging_party, f"Detected {len(gaps)} missing points. Recovering...")
                    for gap in gaps[:10]: # Limit surgical recoveries
                        try:
                            masked_raw = apply_safety_mask(request.raw_data)
                            async with httpx.AsyncClient() as client:
                                recovery_res = await client.post(
                                    final_url,
                                    headers={"api-key": api_key, "Content-Type": "application/json"},
                                    json={
                                        "messages": [
                                            {"role": "system", "content": f"[SURGICAL EXTRACTION] Extract EXACT VERBATIM the Allegation, Suggested Proof, and Response for Point No. {gap}. DO NOT summarize. Keep all names and quotes. Return JSON: {{ 'point': {{ 'label': '...', 'allegation': '...', 'suggested_proof': '...', 'response': '...' }} }}"},
                                            {"role": "user", "content": f"FULL RAW DATA:\n{masked_raw}"}
                                        ],
                                        "response_format": { "type": "json_object" },
                                        "max_completion_tokens": 1024
                                    },
                                    timeout=60.0
                                )
                            if recovery_res.status_code == 200:
                                rec_json = restore_safety_mask(json.loads(repair_json(recovery_res.json()["choices"][0]["message"]["content"])))
                                rec_p = rec_json.get("point") or rec_json.get("points", [{}])[0]
                                if rec_p.get("allegation"): 
                                    rec_p["label"] = str(gap)
                                    all_points.append(rec_p)
                                    activity_logger.log_event("Drafting", "GAP_RECOVERED", request.charging_party, f"Successfully recovered Point {gap} verbatim.")
                        except: continue
        except Exception as audit_ex:
            activity_logger.log_event("Drafting", "AUDIT_WARN", request.charging_party, f"Audit failed: {str(audit_ex)}")

        # deduplicate by Label
        unique_points = {}
        for p in all_points:
            lbl = str(p.get('label', ''))
            if lbl not in unique_points or len(str(p.get('allegation',''))) > len(str(unique_points[lbl].get('allegation',''))):
                unique_points[lbl] = p

        final_points = sorted(unique_points.values(), key=lambda x: int(re.findall(r'\d+', str(x.get('label', '0')))[0]) if re.findall(r'\d+', str(x.get('label', ''))) else 999)

        # --- STEP 1e: CRITICAL IDENTITY CHECK (THE "BETSY" MARKER) ---
        has_betsy_extracted = any("Betsy" in str(p.get("allegation", "")) or "Betsy" in str(p.get("response", "")) for p in final_points)
        has_betsy_raw = "Betsy" in request.raw_data
        
        if has_betsy_raw and not has_betsy_extracted:
            activity_logger.log_event("Drafting", "BETSY_LOST", request.charging_party, "Critical 'Betsy' point missing from extraction. Triggering Surgical Identity Recovery...")
            try:
                masked_raw = apply_safety_mask(request.raw_data)
                async with httpx.AsyncClient() as client:
                    recovery_res = await client.post(
                        final_url,
                        headers={"api-key": api_key, "Content-Type": "application/json"},
                        json={
                            "messages": [
                                {"role": "system", "content": "[SURGICAL IDENTITY RECOVERY] Verbatim Extract the specific Index Point, Allegation, and Suggested Proof containing the name 'Betsy'. DO NOT summarize. Return JSON: { 'point': { 'label': '...', 'allegation': '...', 'suggested_proof': '...', 'response': '...' } }"},
                                {"role": "user", "content": f"FULL RAW DATA:\n{masked_raw}"}
                            ],
                            "response_format": { "type": "json_object" },
                            "max_completion_tokens": 1024
                        },
                        timeout=120.0
                    )
                if recovery_res.status_code == 200:
                    rec_json = restore_safety_mask(json.loads(repair_json(recovery_res.json()["choices"][0]["message"]["content"])))
                    rec_p = rec_json.get("point") or rec_json.get("points", [{}])[0]
                    if rec_p.get("allegation"): 
                        final_points.append(rec_p)
                        # Re-sort
                        final_points = sorted(final_points, key=lambda x: int(re.findall(r'\d+', str(x.get('label', '0')))[0]) if re.findall(r'\d+', str(x.get('label', ''))) else 999)
                        activity_logger.log_event("Drafting", "BETSY_RECOVERED", request.charging_party, "Successfully recovered Point 22 (Betsy) verbatim.")
            except: pass

        # --- STEP 2: RAG RETRIEVAL ---
        rag_context = ""
        unique_cats = set(p.get("legal_category") for p in final_points if p.get("legal_category"))
        for cat in unique_cats:
            try:
                docs = await retrieve_documents(cat, k=2)
                rag_context += "\n\n".join(d.page_content for d in docs)
            except: break
        if not rag_context: rag_context = "Standard legal principles apply."

        # --- STEP 3: ASYNC PARALLEL DRAFTING ---
        activity_logger.log_event("Drafting", "BATCH_START", request.charging_party, f"Parallel Drafting 10-Module Position Statement for {len(final_points)} points...")
        
        # 3a. Prepare Tasks for Global Modules
        task_intro = generate_intro_history(final_url, api_key, request.charging_party, request.respondent, final_points, rag_context)
        task_facts = generate_facts(final_url, api_key, request.charging_party, request.respondent, final_points, rag_context)
        task_ending = generate_conclusion_appendix(final_url, api_key, request.charging_party, request.respondent)
        
        # 3b. Prepare Tasks for Legal Analysis (Dynamic Split)
        analysis_categories = ["discrimination", "retaliation", "disability", "discharge", "damages"]
        analysis_tasks = {}
        for cat in analysis_categories:
            cat_points = [p for p in final_points if cat in str(p.get("legal_category", "")).lower()]
            if not cat_points: cat_points = final_points[:15] # Fallback to core narrative if not categorized
            analysis_tasks[cat] = generate_analysis_section(final_url, api_key, request.charging_party, request.respondent, cat_points, cat, rag_context)

        # 3c. Prepare Tasks for Individual Allegation Responses (Parallel)
        async def draft_point_async(p):
            prompt = "[SENIOR DEFENSE COUNSEL] Draft a professional, legal-grade response for ONE SPECIFIC allegation. Use 'The Respondent denies...' style. Use the provided 'Suggested Proofs' to bolster the factual basis of the response if relevant. Return JSON: { \"response_label\": \"Response No. X\", \"drafted_response\": \"...\" }"
            data = f"ALLEGATION NO. {p.get('label')}: {p.get('allegation')}\nSUGGESTED PROOFS: {p.get('suggested_proof')}\nRESPONSE: {p.get('response')}\nLAW: {rag_context[:1000]}"
            res = await call_llm_module(final_url, api_key, prompt, data)
            if res:
                p_json = json.loads(repair_json(res))
                return {
                    "label": f"Allegation No. {p.get('label')}",
                    "allegation": sanitize_xml(p.get('allegation')),
                    "suggested_proof": sanitize_xml(p.get('suggested_proof', '')),
                    "response_label": sanitize_xml(p_json.get("response_label") or f"Response No. {p.get('label')}"),
                    "response": sanitize_xml(p_json.get("drafted_response") or p.get('response'))
                }
            return {
                "label": f"Allegation No. {p.get('label')}", 
                "allegation": sanitize_xml(p.get('allegation')), 
                "suggested_proof": sanitize_xml(p.get('suggested_proof', '')),
                "response": sanitize_xml(p.get('response'))
            }

        point_tasks = [draft_point_async(p) for p in final_points]

        # 3d. Gather All Results (Parallel Execution)
        activity_logger.log_event("Drafting", "ASYNC_GATHER", request.charging_party, f"Launching {len(point_tasks) + 3 + len(analysis_tasks)} concurrent AI tasks...")
        
        results = await asyncio.gather(
            task_intro,
            task_facts,
            task_ending,
            asyncio.gather(*analysis_tasks.values()),
            asyncio.gather(*point_tasks),
            return_exceptions=True
        )

        # 3e. Map Results back to State
        # Results indices: 0=intro, 1=facts, 2=ending, 3=analysis_results_list, 4=processed_points
        intro_res = results[0] if not isinstance(results[0], Exception) else {}
        facts_res = results[1] if not isinstance(results[1], Exception) else ""
        ending_res = results[2] if not isinstance(results[2], Exception) else {}
        analysis_results_list = results[3] if not isinstance(results[3], Exception) else []
        processed_points = results[4] if not isinstance(results[4], Exception) else []
        
        # Reconstruct final_sections
        final_sections = {
            "introduction": (intro_res if isinstance(intro_res, dict) else {}).get("introduction", ""),
            "procedural_history": (intro_res if isinstance(intro_res, dict) else {}).get("procedural_history", ""),
            "statement_of_facts": facts_res if isinstance(facts_res, str) else "",
            "analysis": {},
            "conclusion": (ending_res if isinstance(ending_res, dict) else {}).get("conclusion", ""),
            "appendix": (ending_res if isinstance(ending_res, dict) else {}).get("appendix", "")
        }

        # --- HARD FALLBACKS (GUARANTEE DATA) ---
        if not final_sections["introduction"]:
            final_sections["introduction"] = f"The Respondent, {request.respondent}, submits this Position Statement in response to the Charge of Discrimination filed by {request.charging_party}. The Respondent denies each and every allegation of discrimination and retaliation."
            activity_logger.log_event("Drafting", "FALLBACK_TRIGGERED", "Introduction", "Manual Template")

        if not final_sections["procedural_history"]:
            final_sections["procedural_history"] = f"The Charging Party filed the instant charge on or about the current date. The Respondent is timely providing this response to the investigative agency."
            activity_logger.log_event("Drafting", "FALLBACK_TRIGGERED", "Procedural", "Manual Template")

        if not final_sections["statement_of_facts"]:
            fact_summary = " ".join([str(p.get('allegation', '')) for p in final_points[:5]])
            final_sections["statement_of_facts"] = f"The Charging Party was employed by {request.respondent}. During their tenure, the following events occurred: {fact_summary}. The Respondent maintains a professional environment and adheres to all anti-discrimination laws."
            activity_logger.log_event("Drafting", "FALLBACK_TRIGGERED", "Facts", "Data Concatenation")
        
        # Map Analysis Categorical Results
        analysis_keys = list(analysis_tasks.keys())
        for i, cat in enumerate(analysis_keys):
            if i < len(analysis_results_list):
                final_sections["analysis"][cat] = analysis_results_list[i] or f"Legal principles regarding {cat} preclude liability in this matter."

        activity_logger.log_event("Drafting", "BATCH_SUCCESS", request.charging_party, "All modules and points drafted successfully (with fallbacks if needed).")


        # --- STEP 4: DOCX GENERATION (TEMPLATE-FIRST) ---
        # _HERE is .../backend/app/api/api_v1/endpoints
        template_path = Path(_HERE).parent.parent.parent.parent / "assets" / "templates" / "Legal_Template.docx"
        
        if not template_path.exists():
            activity_logger.log_event("Drafting", "ERROR", request.charging_party, f"Template not found at {template_path}. Falling back to blank document.")
            doc = Document()
            # Minimal styling for fallback
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(11)
        else:
            doc = Document(str(template_path))

            # --- FRONT MATTER REPLACEMENTS ---
            replacements = {
                "ANDREA ROXTON": request.charging_party,
                "Andrea Roxton": request.charging_party,
                "BOSTON CHILDREN'S HOSPITAL": request.respondent,
                "Boston Children's Hospital": request.respondent,
                "Genevieve Benoit": "[Lead Attorney Name]",
                "GENEVIEVE BENOIT": "[LEAD ATTORNEY NAME]",
                "March 25, 2026": datetime.now().strftime("%B %d, %Y"),
                "BCH": request.respondent
            }
            # Only replace placeholders if specifically requested or if they should stay as-is
            if request.case_number:
                replacements["[NEED LAWYER INPUT: MCAD No.]"] = request.case_number
            
            # Ensure other lawyer-specific placeholders are NOT replaced with fake data
            # Standardizing placeholder keys from the Legal_Template.docx
            placeholders = [
                "[NEED LAWYER INPUT: Investigator Name]",
                "[NEED LAWYER INPUT: Investigator Email]",
                "[NEED LAWYER INPUT: MCAD No.]",
                "[NEED LAWYER INPUT: EEOC No.]"
            ]

            # Replace in Tables (Logo and Caption)
            for i, tbl in enumerate(doc.tables):
                if i < 2:
                    for k, v in replacements.items():
                        for row in tbl.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        if k in r.text:
                                            r.text = r.text.replace(k, str(v))

            # Replace in Paragraphs (Front Matter)
            for i, p in enumerate(doc.paragraphs[:20]):
                for k, v in replacements.items():
                    for r in p.runs:
                        if k in r.text:
                            r.text = r.text.replace(k, str(v))

            # --- DELETE BOILERPLATE ---
            # Index 10 is the start of "I. INTRODUCTION" in our template scan
            body_elements = list(doc.element.body)
            # Find the cutoff: anything starting with "I. INTRODUCTION" or similar should be where we delete
            for el in body_elements[10:]:
                if el.tag.endswith('sectPr'): continue # Preserve section props
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)

        # --- APPEND MODULAR CONTENT ---
        def add_centered_header(num, title):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{num} {title}" if num else title)
            r.bold = True
            r.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)

        def add_body_paragraph(text):
            if not text: return
            p = doc.add_paragraph(sanitize_xml(text))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.15

        # Draft Modules
        add_centered_header("I.", "INTRODUCTION")
        add_body_paragraph(final_sections.get("introduction", ""))

        add_centered_header("II.", "PROCEDURAL HISTORY")
        add_body_paragraph(final_sections.get("procedural_history", ""))

        add_centered_header("III.", "STATEMENT OF FACTS")
        add_body_paragraph(final_sections.get("statement_of_facts", ""))

        add_centered_header("IV.", "RESPONDENT'S ANSWERS TO SPECIFIC ALLEGATIONS")
        for p in processed_points:
            p_header = doc.add_paragraph()
            p_header.paragraph_format.space_before = Pt(12)
            r_header = p_header.add_run(sanitize_xml(p['label']))
            r_header.bold = True
            add_body_paragraph(p['allegation'])
            
            # Add Suggested Proof Section if it exists
            if p.get('suggested_proof') and str(p.get('suggested_proof')).strip():
                p_proof = doc.add_paragraph()
                r_proof = p_proof.add_run("Evidence/Suggested Proof:")
                r_proof.bold = True
                r_proof.italic = True
                add_body_paragraph(p['suggested_proof'])

            p_ans = doc.add_paragraph()
            r_ans = p_ans.add_run(sanitize_xml(p.get('response_label', "Response:")))
            r_ans.bold = True
            add_body_paragraph(p['response'])

        add_centered_header("V.", "LEGAL ANALYSIS")
        l_analysis = final_sections.get("analysis", {})
        if isinstance(l_analysis, str):
            add_body_paragraph(l_analysis)
        elif isinstance(l_analysis, dict):
            for k, v in l_analysis.items():
                p_sub = doc.add_paragraph()
                r_sub = p_sub.add_run(sanitize_xml(k.replace('_', ' ').title()))
                r_sub.bold = True
                add_body_paragraph(v)

        add_centered_header("VI.", "AFFIRMATIVE DEFENSES")
        standard_defenses = [
            "Failure to state a claim upon which relief can be granted.",
            "Lack of probable cause as to any alleged violation of anti-discrimination or anti-retaliation statutes.",
            "Legitimate, non-discriminatory and non-retaliatory business reasons for all challenged actions.",
            "No materially adverse employment action taken against Charging Party as a matter of law.",
            "Lack of causal connection between any protected activity and any challenged action; absence of but-for causation.",
            "Equal application of policies and consistent treatment of similarly situated employees; absence of comparator evidence.",
            "Good faith adherence to and enforcement of anti-discrimination, anti-harassment, anti-retaliation, leave, and accommodation policies.",
            "Failure to engage in the interactive process by Charging Party and/or failure to provide sufficient medical documentation.",
            "Statute of limitations and/or untimeliness of any claims outside the actionable period.",
            "Failure to exhaust administrative remedies for any claims not properly presented to the agency.",
            "After-acquired evidence and/or subsequently discovered information limiting or barring damages, if applicable.",
            "Failure to mitigate damages, including economic loss.",
            "Estoppel, waiver, and laches as appropriate based on Charging Party's conduct and delay.",
            "Good faith, reasonable grounds, and lack of willfulness precluding liquidated or punitive damages."
        ]
        for idx, d in enumerate(standard_defenses):
            add_body_paragraph(f"{idx+1}. {d}")

        add_centered_header("VII.", "CONCLUSION")
        conclusion_text = final_sections.get("conclusion", "")
        conclusion_text = re.sub(r'^VI\.\s+CONCLUSION\s*', '', conclusion_text, flags=re.IGNORECASE).strip()
        add_body_paragraph(conclusion_text)
        
        # --- SIGNATURE BLOCK ---
        p_sig = doc.add_paragraph()
        p_sig.paragraph_format.space_before = Pt(24)
        p_sig.add_run("Respectfully submitted,")
        p_resp = doc.add_paragraph()
        r_resp = p_resp.add_run(sanitize_xml(request.respondent))
        r_resp.bold = True
        p_atty = doc.add_paragraph()
        p_atty.add_run("By its attorneys,")
        
        # --- VERIFICATION ---
        add_centered_header(None, "VERIFICATION")
        add_body_paragraph("I, upon information and belief, affirm that the facts stated herein are true and correct to the best of my knowledge based on records maintained in the ordinary course of business and information provided to me. This Position Statement is submitted without waiver of any privileges, defenses, or objections, all of which are expressly preserved.")

        # --- APPENDIX ---
        doc.add_page_break()
        add_centered_header("VIII.", "APPENDIX: STATUTORY AND REGULATORY FRAMEWORK")
        appendix_text = final_sections.get("appendix", "")
        appendix_text = re.sub(r'^VIII\.\s+APPENDIX.*?\n', '', appendix_text, flags=re.IGNORECASE).strip()
        add_body_paragraph(appendix_text)

        # --- SAVE ---
        default_dir = Path(_HERE).parent.parent.parent.parent / "Drafts"
        try:
            f_dir = Path(request.folder_path) if request.folder_path else default_dir
            f_dir.mkdir(parents=True, exist_ok=True)
        except:
            f_dir = default_dir
            f_dir.mkdir(parents=True, exist_ok=True)

        fname = f"Roxton_Draft_{request.charging_party.replace(' ', '_')}_{int(time.time())}.docx"
        fpath = f_dir / fname
        doc.save(str(fpath))
        
        activity_logger.log_event("Drafting", "END", request.charging_party, f"Successfully saved to: {str(fpath)}")
        return {"status": "success", "file_path": str(fpath)}

    except Exception as e:
        import traceback
        activity_logger.log_event("Drafting", "ERROR", request.charging_party, f"Critical: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

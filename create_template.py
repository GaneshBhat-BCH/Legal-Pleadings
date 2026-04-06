from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

def create_legal_template():
    doc = Document()
    
    # Paths to logos
    # Script is in root. backend/assets/ is the target.
    _ASSETS_DIR = Path("backend/assets")
    LEFT_LOGO = _ASSETS_DIR / "bch_logo.png"
    RIGHT_LOGO = _ASSETS_DIR / "hms_logo.png"

    # 1. Logo Table (Table 0)
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(4.0)
    table.columns[1].width = Inches(2.5)
    
    cell_l = table.cell(0, 0)
    cell_r = table.cell(0, 1)
    
    # Add Logos if found
    if LEFT_LOGO.exists():
        cell_l.paragraphs[0].add_run().add_picture(str(LEFT_LOGO), width=Inches(2.5))
    else:
        cell_l.paragraphs[0].text = "[BCH LOGO MISSING]"
        
    if RIGHT_LOGO.exists():
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.add_run().add_picture(str(RIGHT_LOGO), width=Inches(1.2))
    else:
        cell_r.paragraphs[0].text = "[HMS LOGO MISSING]"

    # 2. Add Standard Header Paragraphs
    doc.add_paragraph("\n")
    
    p_date = doc.add_paragraph("[Current Date]")
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph("\nVIA CERTIFIED MAIL AND EMAIL")
    
    doc.add_paragraph("Legal Division\nMassachusetts Commission Against Discrimination\nOne Ashburton Place, 6th Floor\nBoston, MA 02108")
    
    p_re = doc.add_paragraph("\nRE: [MCAD_DOCKET_NUMBER]")
    p_re.runs[0].bold = True
    
    p_cp = doc.add_paragraph("Charging Party: [CHARGING_PARTY]")
    p_cp.runs[0].bold = True
    
    p_resp = doc.add_paragraph("Respondent: [RESPONDENT]")
    p_resp.runs[0].bold = True
    
    doc.add_paragraph("\nDear Investigator,")
    
    doc.add_paragraph("\nRespondent, Boston Children's Hospital (\"Respondent\" or \"BCH\"), submits this Position Statement in response to the Charge of Discrimination filed by the Charging Party, [CHARGING_PARTY] (\"Charging Party\").")
    
    doc.add_paragraph("\nI. INTRODUCTION")
    doc.paragraphs[-1].runs[0].bold = True
    
    doc.add_paragraph("Respondent is a world-renowned pediatric hospital dedicated to the highest standards of patient care and professional conduct. Respondent denies any and all allegations of discrimination or retaliation and maintains that all employment actions taken with respect to the Charging Party were based on legitimate, non-discriminatory, and non-retaliatory business and operational needs.")
    
    doc.add_paragraph("\nII. RESPONDENT'S BACKGROUND")
    doc.paragraphs[-1].runs[0].bold = True
    
    doc.add_paragraph("Boston Children's Hospital provides complex pediatric specialty care and is a primary teaching hospital of Harvard Medical School. The Hospital maintains strict policies prohibiting discrimination, harassment, and retaliation in any form.")

    doc.add_paragraph("\nIII. COMPLAINT'S ALLEGATIONS")
    doc.paragraphs[-1].runs[0].bold = True
    
    doc.add_paragraph("[The point-by-point drafting will start after this section separator]")

    # Save to the templates folder
    output_path = Path("backend/assets/templates/Legal_Template.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created template with images at: {output_path.absolute()}")

if __name__ == "__main__":
    create_legal_template()
